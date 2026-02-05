import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

def create_word_document():
    # Create a new document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Software Requirements Specification (SRS) - Kings Builder Real Estate Management System"
    doc.core_properties.subject = "Real Estate Management System"
    doc.core_properties.creator = "Real Estate Management Team"
    doc.core_properties.description = "SRS Document for Kings Builder Software"

    # Add title page
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Software Requirements Specification (SRS)")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run("Kings Builder - Real Estate Management System")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Empty paragraph for spacing

    # Add document info
    info_para = doc.add_paragraph()
    info_para.add_run("Document Version: 1.0\n")
    info_para.add_run(f"Date: January 27, 2026\n")
    info_para.add_run("Project: Kings Builder - Property Management System\n")
    info_para.add_run("Prepared by: Real Estate Management Team")
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Add Table of Contents placeholder
    toc_header = doc.add_paragraph()
    toc_header_run = toc_header.add_run("Table of Contents")
    toc_header_run.bold = True
    toc_header_run.font.size = Pt(16)

    doc.add_paragraph()  # Empty paragraph for spacing

    # Read the markdown content
    with open('Kings_Builder_SRS_Document.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Split content by sections
    sections = re.split(r'^##\s+', content, flags=re.MULTILINE)[1:]  # Skip the first split

    for section in sections:
        lines = section.strip().split('\n')
        if not lines[0]:
            continue

        # Extract section title
        section_title = lines[0].strip('# ')

        # Add section header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(section_title)
        header_run.bold = True
        header_run.font.size = Pt(14)

        # Process content for this section
        content_lines = lines[1:]
        current_paragraph = doc.add_paragraph()

        for line in content_lines:
            line = line.rstrip()

            # Check for various markdown elements
            if line.startswith('### '):
                # Subsection
                sub_title = line.strip('# ')
                sub_para = doc.add_paragraph()
                sub_run = sub_para.add_run(sub_title)
                sub_run.bold = True
                sub_run.font.size = Pt(12)
            elif line.startswith('#### '):
                # Sub-subsection
                sub_sub_title = line.strip('# ')
                sub_sub_para = doc.add_paragraph()
                sub_sub_run = sub_sub_para.add_run(sub_sub_title)
                sub_sub_run.bold = True
                sub_sub_run.font.size = Pt(11)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet points
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(line[2:])  # Remove '- ' or '* '
            elif line.startswith('|'):
                # Table
                if line.count('|') >= 2:  # Proper table row
                    parts = line.split('|')
                    if len(parts) > 2:  # Actual table data
                        # For now, just add as plain text - would need more complex parsing for proper tables
                        table_para = doc.add_paragraph()
                        table_para.add_run(line)
            elif line.strip() == '':
                # Empty line - add paragraph break
                current_paragraph = doc.add_paragraph()
            elif line.startswith('```'):
                # Code block
                code_para = doc.add_paragraph()
                # Find the matching closing ```
                code_content = []
                in_code_block = True
                idx = content_lines.index(line)
                for next_line in content_lines[idx+1:]:
                    if next_line.strip().startswith('```'):
                        break
                    code_content.append(next_line)

                run = code_para.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                # Apply shading to distinguish code blocks
                paragraph_format = code_para.paragraph_format
                paragraph_format.left_indent = Inches(0.2)
                paragraph_format.right_indent = Inches(0.2)
            else:
                # Regular paragraph
                if current_paragraph.text:
                    current_paragraph.add_run('\n' + line)
                else:
                    current_paragraph.add_run(line)

    # Save the document
    doc.save('Kings_Builder_SRS_Document.docx')
    print("Word document created successfully: Kings_Builder_SRS_Document.docx")

if __name__ == "__main__":
    # Install python-docx if not already installed
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    create_word_document()